import os
import requests
import re
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from mcp.server.fastmcp import FastMCP
from starlette.applications import Starlette
from starlette.routing import Mount
import uvicorn

# --- RAG IMPORTS ---
from dotenv import load_dotenv
from ddgs import DDGS
from pinecone import Pinecone, ServerlessSpec
from sentence_transformers import SentenceTransformer

def log(msg):
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}")

# ==========================================
# 1. GLOBAL SETUP & AI INITIALIZATION
# ==========================================
load_dotenv()
PINECONE_API_KEY = os.getenv("pinecone_key")

if not PINECONE_API_KEY:
    raise ValueError("🚨 Missing PINECONE_API_KEY! Check your .env file.")

log("⏳ Loading local embedding model (this takes a few seconds)...")
model = SentenceTransformer('all-MiniLM-L6-v2') 

log("🔌 Connecting to Pinecone...")
pc = Pinecone(api_key=PINECONE_API_KEY)
INDEX_NAME = "web-research-test"

if INDEX_NAME not in pc.list_indexes().names():
    log(f"🏗️  Creating new Pinecone index '{INDEX_NAME}'...")
    pc.create_index(
        name=INDEX_NAME, dimension=384, metric="cosine", 
        spec=ServerlessSpec(cloud="aws", region="us-east-1")
    )

index = pc.Index(INDEX_NAME)

# Initialize the MCP Server
mcp = FastMCP("ContentServer")


# ==========================================
# 2. HELPER FUNCTIONS
# ==========================================
def clean_scraped_text(markdown_text):
    text = markdown_text
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    text = re.sub(r'http[s]?://\S+', '', text)
    text = re.sub(r'\[\d+\]', '', text)
    text = re.sub(r'#+\s+', '', text)
    text = re.sub(r'[*^]', '', text)
    text = re.sub(r'-\s\[x\]', '', text)
    text = re.sub(r'_(.*?)_', r'\1', text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n\s*\n', '\n', text).strip()
    return text

def chunk_text(text, chunk_size=250, overlap=50):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk: chunks.append(chunk)
    return chunks

def clean_mermaid(text_from_ai):
    match = re.search(r'```(?:mermaid)?\n(.*?)\n```', text_from_ai, re.DOTALL | re.IGNORECASE)
    if match: code = match.group(1)
    else: code = text_from_ai.replace("```mermaid", "").replace("```", "")
    lines = [line.strip() for line in code.split("\n") if line.strip()]
    return "\n".join(lines)

def fetch_diagram_bytes(mermaid_code):
    cleaned_code = clean_mermaid(mermaid_code)
    log("🛠️  Sending Mermaid code to ChartQuery API...")
    url = "https://api.chartquery.com/v1/diagram"
    payload = {"diagram_type": "mermaid", "diagram_source": cleaned_code, "output_format": "png", "share": True}
    try:
        response = requests.post(url, json=payload, timeout=30)
        if response.status_code == 200:
            render_url = response.json().get("render_url")
            if not render_url: 
                log("❌ ERROR: API succeeded but 'render_url' was missing.")
                return None
            log(f"✅ Render URL received: {render_url}")
            img_response = requests.get(render_url, timeout=30)
            if img_response.status_code == 200 and 'image' in img_response.headers.get('Content-Type', ''):
                log("✅ Diagram successfully downloaded.")
                return img_response.content
            else: 
                log(f"❌ ERROR: Failed to download image. Status: {img_response.status_code}")
                return None
        else: 
            log(f"❌ API ERROR ({response.status_code}): {response.text}")
            return None
    except Exception as e:
        log(f"❌ FATAL REQUEST ERROR: {str(e)}")
        return None

def create_or_edit_word(file_path, content=None, title=None, mermaid_code=None):
    try:
        if not file_path.lower().endswith('.docx'):
            file_path = os.path.splitext(file_path)[0] + ".docx"

        doc = Document(file_path) if os.path.exists(file_path) else Document()
        
        if title:
            t = doc.add_heading(title, 0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if mermaid_code:
            doc.add_heading("System Architecture Diagram", level=2)
            img_bytes = fetch_diagram_bytes(mermaid_code)
            if img_bytes:
                doc.add_picture(BytesIO(img_bytes), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.save(file_path)
                return "Error: Diagram API failed. Check the Server Terminal for the exact syntax error."

        if content and isinstance(content, str) and content.strip():
            blocks = content.split('\n\n')
            for block in blocks:
                block = block.strip()
                if not block: continue

                # Table Logic
                if "|" in block and "\n" in block:
                    lines = block.split('\n')
                    if len([l for l in lines if '|' in l]) > 1:
                        rows = [l.split('|') for l in lines if '|' in l]
                        table_data = []
                        for r in rows:
                            cleaned_row = [cell.strip() for cell in r]
                            if cleaned_row and cleaned_row[0] == '': cleaned_row.pop(0)
                            if cleaned_row and cleaned_row[-1] == '': cleaned_row.pop()
                            if not cleaned_row: continue
                            if not all(all(char in "-: " for char in cell) for cell in cleaned_row):
                                table_data.append(cleaned_row)
                        
                        if table_data:
                            max_cols = max(len(row) for row in table_data)
                            table = doc.add_table(rows=len(table_data), cols=max_cols)
                            table.style = 'Table Grid'
                            for r_idx, row_data in enumerate(table_data):
                                for c_idx, val in enumerate(row_data):
                                    if c_idx < max_cols:
                                        p = table.cell(r_idx, c_idx).paragraphs[0]
                                        for i, part in enumerate(val.split("**")):
                                            run = p.add_run(part)
                                            if i % 2 != 0: run.bold = True
                                        if r_idx == 0:
                                            for run in p.runs: run.bold = True
                            continue

                # Standard Text Logic
                for line in block.split('\n'):
                    line = line.strip()
                    if not line: continue
                    
                    line_lower = line.lower()
                    if ("diagram" in line_lower and "inserted" in line_lower) or line.startswith("*(") or line == "*":
                        continue

                    if line.startswith("**Page ") or line.startswith("Page "):
                        if "Page 1" not in line: doc.add_page_break()
                        clean_line = line.replace("**", "")
                        if ":" in clean_line:
                            clean_title = clean_line.split(":", 1)[1].strip()
                        else:
                            clean_title = re.sub(r'(?i)Page\s+\d+', '', clean_line).strip()
                            
                        if clean_title: doc.add_heading(clean_title, level=1)
                        continue
                        
                    if line.startswith("### "): doc.add_heading(line.replace("### ", "").replace("**", ""), level=3); continue
                    if line.startswith("## "): doc.add_heading(line.replace("## ", "").replace("**", ""), level=2); continue
                    if line.startswith("# "): doc.add_heading(line.replace("# ", "").replace("**", ""), level=1); continue

                    if line.startswith("* ") or line.startswith("- "):
                        p = doc.add_paragraph(style='List Bullet')
                        line = line[2:].strip()
                    else:
                        p = doc.add_paragraph()

                    for i, part in enumerate(line.split("**")):
                        run = p.add_run(part)
                        if i % 2 != 0: run.bold = True

        doc.save(file_path)
        log(f"✅ Document successfully saved to {os.path.abspath(file_path)}")
        return f"Success: Word Document updated at {os.path.abspath(file_path)}"
    except Exception as e:
        log(f"❌ Error creating Word document: {str(e)}")
        return f"Error: {str(e)}"

# ==========================================
# 3. AI TOOLS
# ==========================================

@mcp.tool()
def web_research(query: str) -> str:
    """Call this tool FIRST and ALONE to research real-time information from the web before generating documents.
    It scrapes the internet and uses a Vector Database to extract relevant facts."""
    
    clean_query = query.replace("'", "").replace('"', "").strip()
    log(f"🌍 SERVER RECEIVED TOOL CALL: web_research")
    log(f"🔍 Searching DuckDuckGo for: '{clean_query}'...")
    
    search_results = DDGS().text(clean_query, max_results=3)
    if not search_results:
        log("❌ No search results found on DuckDuckGo.")
        return "Error: No search results found on the web."

    full_text = None
    target_url = None
    
    for index_num, result in enumerate(search_results):
        url = result['href']
        log(f"🕷️ Attempting to scrape [{index_num+1}/3]: {url}")
        
        jina_url = f"https://r.jina.ai/{url}"
        response = requests.get(jina_url, headers={"Accept": "text/markdown"}, timeout=20)
        
        if response.status_code == 200:
            text = response.text
            bad_phrases = ["Target URL returned error 403", "requiring CAPTCHA", "Just a moment..."]
            
            if any(phrase in text for phrase in bad_phrases) or len(text) < 300:
                log("⚠️ Site blocked or returned too little text. Skipping...")
                continue 
            
            full_text = clean_scraped_text(text)
            target_url = url
            log(f"✅ Scraping successful! Cleaned text length: {len(full_text)} characters.")
            break 
        else:
            log(f"❌ Scrape failed with HTTP {response.status_code}.")

    if not full_text:
        log("❌ All top search results failed to scrape.")
        return "Error: All top search results blocked the scraper. Cannot retrieve data."
    
    log("🔪 Chunking text...")
    chunks = chunk_text(full_text, chunk_size=250, overlap=50)
    log(f"✅ Created {len(chunks)} chunks. Generating embeddings...")
    
    embeddings = model.encode(chunks).tolist()
    
    vectors_to_upsert = []
    for i, (chunk, vector) in enumerate(zip(chunks, embeddings)):
        vectors_to_upsert.append({
            "id": f"chunk_{i}",
            "values": vector,
            "metadata": {"text": chunk, "source": target_url}
        })
        
    log("🧠 Upserting vectors to Pinecone...")
    index.upsert(vectors=vectors_to_upsert)
    
    log(f"🔎 Querying Pinecone for: '{clean_query}'...")
    question_vector = model.encode([clean_query]).tolist()[0]
    
    db_results = index.query(
        vector=question_vector,
        top_k=3,
        include_metadata=True
    )
    
    log("🎯 Pinecone retrieval complete. Formatting results...")
    extracted_facts = f"--- RESEARCH CONTEXT (Source: {target_url}) ---\n\n"
    for match in db_results['matches']:
        extracted_facts += match['metadata']['text'] + "\n\n"
        
    log("🧹 Cleaning up Pinecone Index...")
    index.delete(delete_all=True)
    
    log("✅ web_research tool finished successfully.")
    return extracted_facts


@mcp.tool()
def write_to_word(file_path: str, content: str, title: str = None) -> str:
    """Create a structured Word document. 
    - Use '**Page X: Title**' for sections. 
    - Use Markdown style tables.
    - ABSOLUTELY NO PLACEHOLDERS: Do NOT write text like '(Diagram will be inserted here)'.
    CRITICAL INSTRUCTION: If the user asked for a Word document AND a diagram, you MUST call this tool AND the 'insert_diagram' tool at the same time! Do not leave the diagram out!"""
    log(f"📝 SERVER RECEIVED TOOL CALL: write_to_word (File: {file_path})")
    return create_or_edit_word(file_path=file_path, content=content, title=title)


@mcp.tool()
def insert_diagram(file_path: str, mermaid_syntax: str) -> str:
    """Use this tool to append an architectural diagram into an EXISTING Word Document. 
    CRITICAL RULES:
    1. NEVER output the mermaid code to the user in text. Only pass it into this tool's argument.
    2. Syntax MUST be flawless. Start with `graph TD`.
    3. YOU MUST QUOTE ALL NODE LABELS to prevent syntax errors. Example: A["This is text (with parentheses)"] --> B["More text"]
    4. NO TEXT ON ARROWS."""
    log(f"📊 SERVER RECEIVED TOOL CALL: insert_diagram (File: {file_path})")
    
    # --- AGENTIC GUARDRAIL ---
    check_path = file_path if file_path.lower().endswith('.docx') else file_path + '.docx'
    if not os.path.exists(check_path):
        error_msg = "ERROR: Document does not exist yet! You MUST call 'write_to_word' FIRST to generate the text content, and then call 'insert_diagram' afterwards."
        log(f"⚠️  GUARDRAIL TRIGGERED: AI tried to insert diagram before writing document.")
        return error_msg
        
    return create_or_edit_word(file_path=file_path, mermaid_code=mermaid_syntax)

@mcp.tool()
def save_diagram_png(file_path: str, mermaid_syntax: str) -> str:
    """Use this tool when the user asks for a flowchart, diagram, or image, but NOT a Word Document.
    CRITICAL RULES:
    1. NEVER ask the user what format they want. Default to using this tool.
    2. Invent a filename if none is provided (e.g., 'system_architecture.png').
    3. NEVER output the mermaid code to the user in conversational text. 
    4. Syntax MUST be flawless. Start with `graph TD`.
    5. YOU MUST QUOTE ALL NODE LABELS to prevent syntax errors. Example: A["This is text (with parentheses)"] --> B["More text"]
    6. NO TEXT ON ARROWS."""
    log(f"🖼️ SERVER RECEIVED TOOL CALL: save_diagram_png (File: {file_path})")
    if not file_path.lower().endswith('.png'): file_path += '.png'
    img_bytes = fetch_diagram_bytes(mermaid_syntax)
    if img_bytes:
        try:
            with open(file_path, 'wb') as f: f.write(img_bytes)
            log(f"✅ Standalone diagram successfully saved to {file_path}")
            return f"Success: Standalone diagram saved to {os.path.abspath(file_path)}"
        except Exception as e:
            log(f"❌ Error writing standalone PNG: {str(e)}")
            return f"Error writing file: {str(e)}"
    return "Error: Failed to fetch diagram from API."

# --- ASGI Server Setup ---
app = Starlette(
    routes=[
        Mount("/", app=mcp.sse_app())
    ]
)

if __name__ == "__main__":
    log("🚀 Starting MCP Server on SSE transport (http://localhost:8000/sse)")
    uvicorn.run("MCP_Server:app", host="127.0.0.1", port=8000, reload=True)