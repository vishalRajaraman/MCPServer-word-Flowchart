import os
import requests
import re
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from mcp.server.fastmcp import FastMCP
from starlette.applications import Starlette
from starlette.routing import Mount
import uvicorn

# Initialize the MCP Server
mcp = FastMCP("ContentServer")

def clean_mermaid(text_from_ai):
    """Extracts Mermaid code even if the AI surrounds it with conversational text."""
    match = re.search(r'```(?:mermaid)?\n(.*?)\n```', text_from_ai, re.DOTALL | re.IGNORECASE)
    
    if match:
        code = match.group(1)
    else:
        code = text_from_ai.replace("```mermaid", "").replace("```", "")
        
    lines = [line.strip() for line in code.split("\n") if line.strip()]
    return "\n".join(lines)

def fetch_diagram_bytes(mermaid_code):
    """Helper function to fetch the raw PNG bytes from ChartQuery API"""
    cleaned_code = clean_mermaid(mermaid_code)
    
    print("\n" + "="*40)
    print("🛠️  DEBUG: MERMAID CODE SENT TO API:")
    print(cleaned_code)
    print("="*40 + "\n")
    
    url = "https://api.chartquery.com/v1/diagram"
    payload = {
        "diagram_type": "mermaid",
        "diagram_source": cleaned_code,
        "output_format": "png",
        "share": True
    }
    try:
        response = requests.post(url, json=payload, timeout=30)
        
        if response.status_code == 200:
            render_url = response.json().get("render_url")
            if not render_url:
                print("❌ ERROR: API succeeded but 'render_url' was missing in JSON.")
                return None
            print("Render url:",render_url)
            img_response = requests.get(render_url, timeout=30)
            if img_response.status_code == 200 and 'image' in img_response.headers.get('Content-Type', ''):
                print("✅ SUCCESS: Diagram generated and downloaded.")
                return img_response.content
            else:
                print(f"❌ ERROR: Failed to download image from render_url. Status: {img_response.status_code}")
                return None
        else:
            print(f"❌ API ERROR ({response.status_code}): {response.text}")
            return None
            
    except Exception as e:
        print(f"❌ FATAL REQUEST ERROR: {str(e)}")
        return None

def create_or_edit_word(file_path, content=None, title=None, mermaid_code=None):
    try:
        if not file_path.lower().endswith('.docx'):
            file_path = os.path.splitext(file_path)[0] + ".docx"

        doc = Document(file_path) if os.path.exists(file_path) else Document()
        
        if title:
            t = doc.add_heading(title, 0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if mermaid_code:
            doc.add_heading("System Architecture Diagram", level=2)
            img_bytes = fetch_diagram_bytes(mermaid_code)
            if img_bytes:
                doc.add_picture(BytesIO(img_bytes), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.save(file_path)
                return "Error: Diagram API failed. Check the Server Terminal for the exact syntax error."

        if content and isinstance(content, str) and content.strip():
            blocks = content.split('\n\n')
            for block in blocks:
                block = block.strip()
                if not block: continue

                # Table Logic
                if "|" in block and "\n" in block:
                    lines = block.split('\n')
                    if len([l for l in lines if '|' in l]) > 1:
                        rows = [l.split('|') for l in lines if '|' in l]
                        table_data = []
                        for r in rows:
                            cleaned_row = [cell.strip() for cell in r]
                            if cleaned_row and cleaned_row[0] == '': cleaned_row.pop(0)
                            if cleaned_row and cleaned_row[-1] == '': cleaned_row.pop()
                            if not cleaned_row: continue
                            if not all(all(char in "-: " for char in cell) for cell in cleaned_row):
                                table_data.append(cleaned_row)
                        
                        if table_data:
                            max_cols = max(len(row) for row in table_data)
                            table = doc.add_table(rows=len(table_data), cols=max_cols)
                            table.style = 'Table Grid'
                            for r_idx, row_data in enumerate(table_data):
                                for c_idx, val in enumerate(row_data):
                                    if c_idx < max_cols:
                                        p = table.cell(r_idx, c_idx).paragraphs[0]
                                        for i, part in enumerate(val.split("**")):
                                            run = p.add_run(part)
                                            if i % 2 != 0: run.bold = True
                                        if r_idx == 0:
                                            for run in p.runs: run.bold = True
                            continue

                # Standard Text Logic
                for line in block.split('\n'):
                    line = line.strip()
                    if not line: continue
                    
                    line_lower = line.lower()
                    if ("diagram" in line_lower and "inserted" in line_lower) or line.startswith("*(") or line == "*":
                        continue

                    # --- UPDATED PAGE FORMATTING LOGIC ---
                    if line.startswith("**Page ") or line.startswith("Page "):
                        if "Page 1" not in line: 
                            doc.add_page_break()
                        
                        # Remove the bold stars
                        clean_line = line.replace("**", "")
                        
                        # Split by the colon and keep only the actual title
                        if ":" in clean_line:
                            clean_title = clean_line.split(":", 1)[1].strip()
                        else:
                            clean_title = re.sub(r'(?i)Page\s+\d+', '', clean_line).strip()
                            
                        if clean_title:
                            doc.add_heading(clean_title, level=1)
                        continue
                        
                    if line.startswith("### "): doc.add_heading(line.replace("### ", "").replace("**", ""), level=3); continue
                    if line.startswith("## "): doc.add_heading(line.replace("## ", "").replace("**", ""), level=2); continue
                    if line.startswith("# "): doc.add_heading(line.replace("# ", "").replace("**", ""), level=1); continue

                    if line.startswith("* ") or line.startswith("- "):
                        p = doc.add_paragraph(style='List Bullet')
                        line = line[2:].strip()
                    else:
                        p = doc.add_paragraph()

                    for i, part in enumerate(line.split("**")):
                        run = p.add_run(part)
                        if i % 2 != 0: run.bold = True

        doc.save(file_path)
        return f"Success: Word Document updated at {os.path.abspath(file_path)}"
    except Exception as e:
        return f"Error: {str(e)}"

# --- TOOL 1: Create Word Document ---
@mcp.tool()
def write_to_word(file_path: str, content: str, title: str = None) -> str:
    """Create a structured Word document. 
    - Use '**Page X: Title**' for sections. 
    - Use Markdown style tables.
    - ABSOLUTELY NO PLACEHOLDERS: Do NOT write text like '(Diagram will be inserted here)'.
    
    CRITICAL INSTRUCTION: If the user asked for a Word document AND a diagram, you MUST call this tool AND the 'insert_diagram' tool at the same time! Do not leave the diagram out!"""
    return create_or_edit_word(file_path=file_path, content=content, title=title)

# --- TOOL 2: Append Diagram to Existing Word Doc ---
@mcp.tool()
def insert_diagram(file_path: str, mermaid_syntax: str) -> str:
    """Use this tool to append a Mermaid.js architectural diagram image into an EXISTING Word Document. 
    CRITICAL: If the user asked for a new document AND a diagram, you MUST call 'write_to_word' AND this tool together in the same response.
    
    CRITICAL SYNTAX RULES:
    1. The 'mermaid_syntax' argument MUST contain RAW MERMAID CODE ONLY. Do NOT include any conversational English text.
    2. Start the code exactly with `graph TD`.
    3. Quote ALL node labels. Example: A["Client"] --> B["Server"]
    4. NO TEXT ON ARROWS. Do NOT use the pipe `|` character."""
    return create_or_edit_word(file_path=file_path, mermaid_code=mermaid_syntax)

# --- TOOL 3: Save Diagram as Standalone PNG ---
@mcp.tool()
def save_diagram_png(file_path: str, mermaid_syntax: str) -> str:
    """Use this tool when the user ONLY wants an image file (.png), NOT a Word document (.docx).
    Generates an architectural diagram using Mermaid.js and saves it directly as a standalone PNG.
    
    CRITICAL SYNTAX RULES:
    1. The 'mermaid_syntax' argument MUST contain RAW MERMAID CODE ONLY. Do NOT include any conversational English text.
    2. Start the code exactly with `graph TD`.
    3. Quote ALL node labels. Example: A["Client"] --> B["Server"]
    4. NO TEXT ON ARROWS. Do NOT use the pipe `|` character."""
    if not file_path.lower().endswith('.png'):
        file_path += '.png'
        
    img_bytes = fetch_diagram_bytes(mermaid_syntax)
    if img_bytes:
        try:
            with open(file_path, 'wb') as f:
                f.write(img_bytes)
            return f"Success: Standalone diagram saved to {os.path.abspath(file_path)}"
        except Exception as e:
            return f"Error writing file: {str(e)}"
    return "Error: Failed to fetch diagram from API."

# --- ASGI Server Setup ---
app = Starlette(
    routes=[
        Mount("/", app=mcp.sse_app())
    ]
)

if __name__ == "__main__":
    print("🚀 Starting MCP Server on SSE transport (http://localhost:8000/sse)")
    uvicorn.run("wordMCP:app", host="127.0.0.1", port=8000, reload=True)